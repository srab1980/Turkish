#!/usr/bin/env python3
"""
Script to extract content from Turkish A1 curriculum documents
"""

import os
import json
import re
from typing import Dict, List, Any

def extract_docx_content(file_path: str) -> Dict[str, Any]:
    """Extract content from DOCX file using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = {
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content["paragraphs"].append({
                    "text": para.text.strip(),
                    "style": para.style.name if para.style else "Normal"
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            content["tables"].append(table_data)
        
        # Extract metadata
        core_props = doc.core_properties
        content["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        return content
        
    except ImportError:
        print("python-docx not installed. Attempting alternative extraction...")
        return extract_docx_alternative(file_path)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return {"error": str(e)}

def extract_docx_alternative(file_path: str) -> Dict[str, Any]:
    """Alternative DOCX extraction using zipfile"""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        content = {
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        
        with zipfile.ZipFile(file_path, 'r') as docx:
            # Extract main document
            try:
                xml_content = docx.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                # Define namespaces
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Extract paragraphs
                for para in root.findall('.//w:p', namespaces):
                    text_elements = para.findall('.//w:t', namespaces)
                    para_text = ''.join([elem.text or '' for elem in text_elements])
                    if para_text.strip():
                        content["paragraphs"].append({
                            "text": para_text.strip(),
                            "style": "Normal"
                        })
                
                # Extract tables
                for table in root.findall('.//w:tbl', namespaces):
                    table_data = []
                    for row in table.findall('.//w:tr', namespaces):
                        row_data = []
                        for cell in row.findall('.//w:tc', namespaces):
                            cell_text_elements = cell.findall('.//w:t', namespaces)
                            cell_text = ''.join([elem.text or '' for elem in cell_text_elements])
                            row_data.append(cell_text.strip())
                        if row_data:
                            table_data.append(row_data)
                    if table_data:
                        content["tables"].append(table_data)
                        
            except Exception as e:
                print(f"Error parsing document XML: {e}")
        
        return content
        
    except Exception as e:
        print(f"Error with alternative extraction: {e}")
        return {"error": str(e)}

def analyze_curriculum_structure(content: Dict[str, Any]) -> Dict[str, Any]:
    """Analyze the curriculum structure and extract learning elements"""
    
    analysis = {
        "units": [],
        "vocabulary": [],
        "grammar_points": [],
        "exercises": [],
        "themes": []
    }
    
    paragraphs = content.get("paragraphs", [])
    tables = content.get("tables", [])
    
    # Patterns to identify different content types
    unit_pattern = re.compile(r'(ünite|unit|bölüm|chapter)\s*(\d+)', re.IGNORECASE)
    lesson_pattern = re.compile(r'(ders|lesson)\s*(\d+)', re.IGNORECASE)
    vocab_pattern = re.compile(r'(kelime|vocabulary|sözcük)', re.IGNORECASE)
    grammar_pattern = re.compile(r'(dilbilgisi|grammar|gramer)', re.IGNORECASE)
    
    current_unit = None
    current_lesson = None
    
    for para in paragraphs:
        text = para["text"]
        
        # Check for unit headers
        unit_match = unit_pattern.search(text)
        if unit_match:
            current_unit = {
                "number": unit_match.group(2),
                "title": text,
                "lessons": [],
                "vocabulary": [],
                "grammar": []
            }
            analysis["units"].append(current_unit)
            continue
        
        # Check for lesson headers
        lesson_match = lesson_pattern.search(text)
        if lesson_match and current_unit:
            current_lesson = {
                "number": lesson_match.group(2),
                "title": text,
                "content": []
            }
            current_unit["lessons"].append(current_lesson)
            continue
        
        # Check for vocabulary sections
        if vocab_pattern.search(text):
            if current_unit:
                current_unit["vocabulary"].append(text)
            analysis["vocabulary"].append(text)
            continue
        
        # Check for grammar sections
        if grammar_pattern.search(text):
            if current_unit:
                current_unit["grammar"].append(text)
            analysis["grammar_points"].append(text)
            continue
        
        # Add to current lesson content
        if current_lesson:
            current_lesson["content"].append(text)
    
    # Analyze tables for structured content
    for table in tables:
        if len(table) > 1:  # Has header and data
            # Check if it's a vocabulary table
            header = table[0] if table else []
            if any(word in ' '.join(header).lower() for word in ['türkçe', 'turkish', 'english', 'ingilizce']):
                for row in table[1:]:
                    if len(row) >= 2:
                        vocab_item = {
                            "turkish": row[0],
                            "english": row[1] if len(row) > 1 else "",
                            "pronunciation": row[2] if len(row) > 2 else ""
                        }
                        analysis["vocabulary"].append(vocab_item)
    
    return analysis

def main():
    """Main function to extract and analyze curriculum documents"""
    
    curriculum_dir = "Curriculum"
    output_dir = "curriculum_analysis"
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    
    documents = [
        ("A1DersKtabi.docx", "textbook"),
        ("A1alimaKtabi.docx", "workbook")
    ]
    
    for doc_name, doc_type in documents:
        file_path = os.path.join(curriculum_dir, doc_name)
        
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
        
        print(f"Extracting content from {doc_name}...")
        
        # Extract content
        content = extract_docx_content(file_path)
        
        if "error" in content:
            print(f"Error extracting {doc_name}: {content['error']}")
            continue
        
        # Analyze structure
        analysis = analyze_curriculum_structure(content)
        
        # Save raw content
        with open(os.path.join(output_dir, f"{doc_type}_content.json"), 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)
        
        # Save analysis
        with open(os.path.join(output_dir, f"{doc_type}_analysis.json"), 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)
        
        print(f"Extracted {len(content.get('paragraphs', []))} paragraphs and {len(content.get('tables', []))} tables")
        print(f"Found {len(analysis.get('units', []))} units and {len(analysis.get('vocabulary', []))} vocabulary items")
        print(f"Results saved to {output_dir}/")

if __name__ == "__main__":
    main()
